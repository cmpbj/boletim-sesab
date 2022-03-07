import re
import requests
from requests.exceptions import HTTPError
from datetime import date
from datetime import datetime

from docx import Document
from docx.shared import Inches
from selenium import webdriver
from selenium.webdriver.common.by import By

"""**Primeiro, vamos criar um crawler que colherá o conteúdo da página desejada para que possamos coletar os dados**

"""

def crawl_website(url: str) -> str:
    
    try:
        pagina = requests.get(url)
        pagina.raise_for_status()
    except HTTPError as exc:
        print(exc)
    else:
        return pagina.text

"""**Utilizamos a função re.compile para compilar um padrão de expressão regular. A ideia é usar essas expressões para correspondência posterior com o que queremos encontrar no Boletim da Sesab**"""

REGEXP_CASOS = re.compile("([0-9.]+) casos confirmados")
REGEXP_MORTES = re.compile("([0-9.]+) tiveram óbito confirmado")
REGEXP_NOVOS_CASOS = re.compile("Nas últimas 24 horas, foram registrados ([0-9.]+) casos de Covid-19")
REGEXP_NOVAS_MORTES = re.compile("e mais ([0-9.]+) óbitos")
REGEXP_RECUPERADOS = re.compile("([0-9.]+) já são considerados recuperados") 
REGEXP_CASOS_ATIVOS = re.compile("([0-9.]+) casos ativos de Covid-19 na Bahia")
REGEXP_CASOS_DESCARTADOS = re.compile("([0-9.]+) casos descartados")
REGEXP_CASOS_INVESTIGADOS = re.compile("([0-9.]+) em investigação")
REGEXP_VACINACAO_UMA_DOSE = re.compile("([0-9.]+) pessoas vacinadas com a primeira dose")
REGEXP_VACINACAO_DUAS_DOSES = re.compile("([0-9.]+) com a segunda dose ou dose única")
REGEXP_VACINACAO_DOSE_REFORCO = re.compile("([0-9.]+) com a dose de reforço")
REGEXP_VACINACAO_CRIANCAS = re.compile("([0-9.]+) crianças já foram imunizadas")
REGEXP_VACINACAO_CRIANCAS_SEGUNDA_DOSE = re.compile("([0-9.]+) já tomaram também a segunda dose")

"""**Agora vamos definir as datas que serão utilizadas para verificar se o boletim publicado pela Sesab é o mais atual e também para escrever nosso texto**"""

DIAS = [
    'segunda',
    'terça',
    'quarta',
    'quinta',
    'sexta',
    'sábado',
    'domingo'
]

#Utilizaremos essa data para verificar se o boletim é o mais atual
data_atual = datetime.now()
data_em_texto = data_atual.strftime('%d/%m/%Y')

#Aqui, precisamos do dia da semana correspondente à data atual para escrever o boletim
indice_da_semana = data_atual.weekday()

dia_da_semana_por_extenso = DIAS[indice_da_semana]

if dia_da_semana_por_extenso == 'sábado' or dia_da_semana_por_extenso == 'domingo':
    dia = 'neste ' + dia_da_semana_por_extenso
else:
    dia = 'nesta ' + dia_da_semana_por_extenso
    
dia_da_semana = data_atual.day    

#Utilizaremos esta ultima data para salvar o documento com a data do dia que ele foi escrito
data_do_arquivo = '{}-{}-{}'.format(data_atual.day, data_atual.month, data_atual.year)

"""**Vamos usar o pacote Selenium para navegar no portal onde são publicados os boletins da Sesab e para encontrar o boletim propriamente dito. Para o Selenium funcionar, é preciso baixar o webdriver correspondente a seu navagedor. No caso do Google Chrome, você pode obter o webdriver através desse [link](https://chromedriver.chromium.org/downloads). É importante que esse arquivo esteja na mesma pasta em que estamos rodando esse script. Em seguinda, após abrirmos a página da Sesab, vamos usar a função find_element para encontrar o boletim. Veja aqui a documentação do [Selenium](https://selenium-python.readthedocs.io/locating-elements.html)**"""

navegador = webdriver.Chrome()
navegador.get('http://www.saude.ba.gov.br/category/emergencias-em-saude/')

titulo = navegador.find_element(by=By.PARTIAL_LINK_TEXT, value='casos ativos')

"""**Depois de encontrar o boletim na página, vamos acessá-lo**"""

titulo.click()

#Vamos ver qual a data de publicação do boletim
data_hora = navegador.find_element(by=By.TAG_NAME, value='strong')

"""
- Se a data do dia for igual a data de publicação, vamos pegar a URL da página e passar o crawler 
- Depois, utilizamos a função findall do pacote re para encontrar os dados.
- Caso a data do dia seja diferente da data da publicação, o programa avisa que o novo 
boletim ainda não saiu e retorna para a página anterior
- Foi preciso transformar as listas encontradas com findall em conjuntos e depois converter 
em lista novamente para evitar dados duplicados
"""

if data_em_texto in data_hora.text:
    print("O boletim do dia saiu")
    get_url = navegador.current_url
    conteudo = crawl_website(url=get_url)
    casos = list(set(REGEXP_CASOS.findall(conteudo)))
    mortes = list(set(REGEXP_MORTES.findall(conteudo)))
    novos_casos = list(set(REGEXP_NOVOS_CASOS.findall(conteudo)))
    novas_mortes = list(set(REGEXP_NOVAS_MORTES.findall(conteudo)))
    recuperados = list(set(REGEXP_RECUPERADOS.findall(conteudo)))
    casos_ativos = list(set(REGEXP_CASOS_ATIVOS.findall(conteudo)))
    casos_descartados = list(set(REGEXP_CASOS_DESCARTADOS.findall(conteudo)))
    casos_investigacao = list(set(REGEXP_CASOS_INVESTIGADOS.findall(conteudo)))
    vacina_uma_dose = list(set(REGEXP_VACINACAO_UMA_DOSE.findall(conteudo)))
    vacina_duas_doses = list(set(REGEXP_VACINACAO_DUAS_DOSES.findall(conteudo)))
    vacina_reforco = list(set(REGEXP_VACINACAO_DOSE_REFORCO.findall(conteudo)))
    vacina_criancas = list(set(REGEXP_VACINACAO_CRIANCAS.findall(conteudo)))
    vacina_criancas_segunda_dose = list(set(REGEXP_VACINACAO_CRIANCAS_SEGUNDA_DOSE.findall(conteudo)))
else:
    print("O boletim novo ainda não saiu")
    navegador.back()

"""**Com os dados em mãos, agora podemos escrever nosso boletim**"""

documento = Document()

documento.add_paragraph(f'Opção de título 01: Bahia registra {novos_casos[0]} novos casos de Covid-19 e {novas_mortes[0]} mortes em 24 horas')

documento.add_paragraph(f'Opção de subtítulo 01: O boletim epidemiológico divulgado pela Sesab {dia} ({dia_da_semana}) mostra que a Bahia atingiu {casos_ativos[0]} casos ativos da doença')

documento.add_paragraph(f'Opção de título 02: Bahia registra {casos_ativos[0]} casos ativos de Covid-19 e {novas_mortes[0]} mortes pela doença')

documento.add_paragraph(f'Opção de subtítulo 02: Ao todo, são {casos[0]} casos confirmados desde o início da pandemia')

documento.add_paragraph(f'A Bahia registrou {novos_casos[0]} novos casos de Covid-19 e {novas_mortes[0]} mortes pela doença nas últimas 24 horas. Segundo o boletim epidemiológico da Secretaria de Saúde do Estado da Bahia (Sesab), divulgado {dia} ({dia_da_semana}), dos {casos[0]} casos confirmados desde o início da pandemia, {recuperados[0]} já são considerados recuperados e {mortes[0]} tiveram morte confirmada.')

documento.add_paragraph(f'No estado, atualmente, são {casos_ativos[0]} casos ativos do novo coronavírus, além de {casos_descartados[0]} descartados e {casos_investigacao[0]} em investigação.')

documento.add_paragraph('Vacinação')

documento.add_paragraph(f'Até o momento temos {vacina_uma_dose[0]} pessoas vacinadas com a primeira dose, {vacina_duas_doses[0]} com a segunda dose ou dose única e {vacina_reforco[0]} com a dose de reforço. Do público de 5 a 11 anos, {vacina_criancas[0]} crianças já foram imunizadas com a primeira dose e {vacina_criancas_segunda_dose[0]} já tomaram a segunda dose do imunizante.')

documento.save(f'boletim {data_do_arquivo}')